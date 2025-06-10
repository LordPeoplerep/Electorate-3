import json
from docx import Document

with open("candidates.json", "r") as f:
    candidates = json.load(f)

doc = Document()
doc.add_heading("Election Results", 0)

for c in candidates:
    doc.add_paragraph(f"Name: {c['name']}")
    doc.add_paragraph(f"Total Votes: {c['total']}")
    doc.add_paragraph(f"Authenticated Votes: {c['auth']}")
    doc.add_paragraph(f"% Authenticated: {round((c['auth'] / c['total']) * 100, 2)}%")
    doc.add_paragraph("")

doc.save("election_results.docx")